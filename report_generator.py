"""
NanoJS Report Generator
=========================
Reads scan_results.json and produces a Word (.docx) disclosure report.
Author: NanoJS Investigations
"""

import json, os, sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY      = RGBColor(0x0A, 0x1F, 0x44)
ACCENT    = RGBColor(0x00, 0x82, 0xCA)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
RED       = RGBColor(0xC0, 0x39, 0x2B)

SEVERITY_COLORS = {
    "CRITICAL": "C0392B", "HIGH": "E67E22",
    "MEDIUM":   "F39C12", "LOW":  "27AE60",
}
SEVERITY_LABELS = {
    "CRITICAL": "CRITICAL", "HIGH": "HIGH",
    "MEDIUM":   "MEDIUM",   "LOW":  "LOW",
}

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_rule(doc, color="0082CA", size=6):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(2)

def heading(doc, text, level=1):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {1: 18, 2: 13, 3: 11}
    colors = {1: NAVY, 2: ACCENT, 3: DARK_GRAY}
    run.font.size      = Pt(sizes.get(level, 11))
    run.font.bold      = True
    run.font.color.rgb = colors.get(level, DARK_GRAY)
    p.paragraph_format.space_before = Pt({1:18,2:12,3:8}.get(level,8))
    p.paragraph_format.space_after  = Pt(4)
    if level == 1:
        add_rule(doc)

def body(doc, text, italic=False, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size   = Pt(10)
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)

def info_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        kc = t.rows[i].cells[0]
        vc = t.rows[i].cells[1]
        set_cell_bg(kc, "0A1F44")
        set_cell_bg(vc, "F4F6F8")
        kr = kc.paragraphs[0].add_run(k)
        kr.font.bold = True; kr.font.color.rgb = WHITE; kr.font.size = Pt(9)
        kc.width = Inches(2.0)
        vr = vc.paragraphs[0].add_run(str(v))
        vr.font.size = Pt(9); vr.font.color.rgb = DARK_GRAY
        vc.width = Inches(4.5)
    doc.add_paragraph()

def code_block(doc, code, max_lines=60):
    lines = code.strip().split("\n")[:max_lines]
    t    = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, "F0F0F0")
    cell.paragraphs[0]._element.clear()
    for line in lines:
        p   = cell.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        run.font.color.rgb = DARK_GRAY
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()

def severity_badge(doc, severity):
    color = SEVERITY_COLORS.get(severity, "999999")
    p   = doc.add_paragraph()
    run = p.add_run(f"  {severity}  ")
    run.font.bold = True; run.font.size = Pt(10); run.font.color.rgb = WHITE
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color)
    rPr.append(shd)
    p.paragraph_format.space_after = Pt(6)

REMEDIATIONS = {
    "REENTRANCY":       "Apply Checks-Effects-Interactions (CEI) pattern. Update state BEFORE external calls. Use OpenZeppelin ReentrancyGuard.",
    "FLASH_LOAN":       "Replace spot price reads with TWAP oracle (Uniswap V3 TWAP or Chainlink). Never use balanceOf(address(this)) for pricing.",
    "ACCESS_CONTROL":   "Add onlyOwner or AccessControl modifiers to all privileged functions. Use OpenZeppelin Ownable or AccessControl.",
    "SELFDESTRUCT":     "Remove selfdestruct if non-essential. If required, gate behind multi-sig or timelock.",
    "INTEGER_OVERFLOW": "Upgrade to Solidity >=0.8.0 for built-in overflow checks. Otherwise use OpenZeppelin SafeMath.",
    "PROXY_UPGRADE":    "Restrict upgradeTo to multi-sig admin or timelock. Use OpenZeppelin TransparentUpgradeableProxy.",
}

def generate_report(results, output_path=None):
    if not results:
        print("No findings to report.")
        return ""

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    chains    = list({r["contract"]["chain"] for r in results})
    all_finds = [f for r in results for f in r["findings"]]
    report_id = f"NanoJS-{datetime.utcnow().strftime('%Y%m%d-%H%M')}"
    date_str  = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    # ── Cover page ──────────────────────────────────────────
    add_rule(doc, "0082CA", 18)
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("SMART CONTRACT VULNERABILITY\nDISCLOSURE REPORT")
    r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p2 = doc.add_paragraph()
    r2 = p2.add_run("NanoJS Investigations — Responsible Disclosure Series")
    r2.font.size = Pt(11); r2.font.italic = True; r2.font.color.rgb = ACCENT
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_rule(doc, "C0392B", 6)
    doc.add_paragraph()

    info_table(doc, [
        ("Report ID",       report_id),
        ("Investigator",    "NanoJS (NanoJS Investigations)"),
        ("Contact",         "nanojs@proton.me | github.com/NanoJS10"),
        ("Date Issued",     date_str),
        ("Chain(s)",        ", ".join(chains)),
        ("Contracts",       str(len(results))),
        ("Total Findings",  str(len(all_finds))),
        ("Classification",  "CONFIDENTIAL — FOR PROJECT TEAM ONLY"),
    ])

    # Disclaimer
    dt = doc.add_table(rows=1, cols=1)
    dc = dt.rows[0].cells[0]
    set_cell_bg(dc, "FEF9E7")
    dp = dc.paragraphs[0]
    dp.add_run("IMPORTANT: ").font.bold = True
    dp.add_run(
        "This report is for responsible disclosure only. PoC contracts are "
        "prototypes demonstrating vulnerability existence. Do NOT deploy against "
        "live contracts without written project team authorisation."
    ).font.size = Pt(9)
    doc.add_page_break()

    # ── Executive Summary ───────────────────────────────────
    heading(doc, "1. Executive Summary", 1)
    counts = {"CRITICAL":0,"HIGH":0,"MEDIUM":0,"LOW":0}
    for r in results:
        for f in r["findings"]:
            counts[f["severity"]] = counts.get(f["severity"],0) + 1
    body(doc,
        f"NanoJS Investigations analysed {len(results)} contract(s) across "
        f"{', '.join(chains)}, identifying {len(all_finds)} vulnerability "
        f"finding(s) via bytecode analysis, source code pattern matching, and ABI inspection."
    )
    heading(doc, "Finding Summary", 2)
    st = doc.add_table(rows=5, cols=3)
    st.style = "Table Grid"
    for i, h in enumerate(["Severity","Count","Risk"]):
        c = st.rows[0].cells[i]
        set_cell_bg(c, "0A1F44")
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9)
    sev_data = [
        ("CRITICAL", counts["CRITICAL"], "Immediate fund loss / takeover"),
        ("HIGH",     counts["HIGH"],     "Significant financial risk"),
        ("MEDIUM",   counts["MEDIUM"],   "Exploitable under conditions"),
        ("LOW",      counts["LOW"],      "Best practice issue"),
    ]
    for i, (s, c, imp) in enumerate(sev_data):
        row = st.rows[i+1]
        set_cell_bg(row.cells[0], SEVERITY_COLORS[s])
        set_cell_bg(row.cells[1], "F4F6F8")
        set_cell_bg(row.cells[2], "F4F6F8")
        r0 = row.cells[0].paragraphs[0].add_run(s)
        r0.font.bold = True; r0.font.color.rgb = WHITE; r0.font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(str(c)).font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run(imp).font.size = Pt(9)
    doc.add_paragraph()

    # ── Detailed Findings ───────────────────────────────────
    heading(doc, "2. Detailed Findings", 1)
    fnum = 1
    for result in results:
        contract = result["contract"]
        heading(doc, f"Contract: {contract['contract_address']}", 2)
        info_table(doc, [
            ("Chain",     contract["chain"]),
            ("Address",   contract["contract_address"]),
            ("Deployer",  contract["deployer"]),
            ("Block",     str(contract["block"])),
            ("Time",      contract["timestamp"]),
            ("Explorer",  contract["explorer_url"]),
            ("Source",    "Verified" if result["source_available"] else "Bytecode only"),
            ("ABI",       "Available" if result["abi_available"] else "Not available"),
        ])
        for f in result["findings"]:
            heading(doc, f"Finding {fnum:02d} — {f['name']}", 3)
            severity_badge(doc, f["severity"])
            info_table(doc, [
                ("Vuln ID",    f["vuln_id"]),
                ("Severity",   f["severity"]),
                ("Confidence", f"{f['confidence']}%"),
                ("CWE",        f["cwe"]),
                ("Reference",  f["reference"]),
            ])
            heading(doc, "Description", 3)
            body(doc, f["description"])
            heading(doc, "Evidence", 3)
            for ev in f["evidence"]:
                p   = doc.add_paragraph(style="List Bullet")
                p.add_run(ev).font.size = Pt(9)
            heading(doc, "Remediation", 3)
            body(doc, REMEDIATIONS.get(f["vuln_id"], "Review and apply appropriate access controls."))
            doc.add_paragraph()
            fnum += 1

    # ── PoC Section ─────────────────────────────────────────
    heading(doc, "3. Proof-of-Concept Prototypes", 1)
    body(doc,
        "The following Solidity contracts demonstrate each vulnerability. "
        "These are disclosure prototypes only — do NOT deploy on mainnet.",
        italic=True, color=RED
    )
    doc.add_paragraph()
    pnum = 1
    for result in results:
        contract = result["contract"]
        for vuln_id, poc_code in result["pocs"].items():
            f = next((x for x in result["findings"] if x["vuln_id"]==vuln_id), {})
            heading(doc, f"PoC {pnum:02d} — {f.get('name', vuln_id)}", 2)
            info_table(doc, [
                ("Target",      contract["contract_address"]),
                ("Chain",       contract["chain"]),
                ("Severity",    f.get("severity","N/A")),
                ("File",        f"poc_{pnum:02d}_{vuln_id.lower()}.sol"),
            ])
            body(doc, "Solidity Prototype (testnet only):")
            code_block(doc, poc_code, max_lines=80)
            pnum += 1

    # ── Disclosure Timeline ─────────────────────────────────
    heading(doc, "4. Disclosure Timeline", 1)
    today = datetime.utcnow().strftime("%Y-%m-%d")
    tl = [
        (today,           "Discovery and scan completed"),
        ("Within 24h",    "Report delivered to project team"),
        ("Day 3",         "Acknowledgement requested"),
        ("Day 14",        "Fix or mitigation confirmation expected"),
        ("Day 90",        "Public disclosure if no response (industry standard)"),
    ]
    tt = doc.add_table(rows=len(tl)+1, cols=2)
    tt.style = "Table Grid"
    for i, h in enumerate(["Milestone","Action"]):
        c = tt.rows[0].cells[i]
        set_cell_bg(c, "0A1F44")
        r = c.paragraphs[0].add_run(h)
        r.font.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9)
    for i, (d, a) in enumerate(tl):
        row = tt.rows[i+1]
        bg  = "F4F6F8" if i%2==0 else "FFFFFF"
        set_cell_bg(row.cells[0], bg); set_cell_bg(row.cells[1], bg)
        row.cells[0].paragraphs[0].add_run(d).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(a).font.size = Pt(9)
    doc.add_paragraph()

    # ── Signature ────────────────────────────────────────────
    heading(doc, "5. Investigator Statement", 1)
    body(doc,
        "This report was produced independently by NanoJS Investigations as part "
        "of an ongoing responsible disclosure research program. No financial "
        "compensation was received prior to this disclosure. All findings are "
        "based on publicly accessible on-chain data and standard static analysis."
    )
    doc.add_paragraph()
    add_rule(doc, "0082CA", 6)
    doc.add_paragraph()
    info_table(doc, [
        ("Handle",      "NanoJS / NanoJS10"),
        ("Organisation","NanoJS Investigations"),
        ("Email",       "nanojs@proton.me"),
        ("GitHub",      "github.com/NanoJS10"),
        ("Portfolio",   "waliusoji.github.io"),
        ("Generated",   date_str),
    ])
    p = doc.add_paragraph()
    r = p.add_run('"Data on the blockchain is permanent — analysis makes it powerful."')
    r.font.italic = True; r.font.color.rgb = ACCENT

    # Save
    if not output_path:
        output_path = os.path.join(os.getcwd(), f"{report_id}_Disclosure_Report.docx")
    doc.save(output_path)
    print(f"\n✅ Report saved: {output_path}")
    return output_path


if __name__ == "__main__":
    results_file = sys.argv[1] if len(sys.argv) > 1 else "scan_results.json"
    if not os.path.exists(results_file):
        print(f"ERROR: {results_file} not found. Run scanner.py first.")
        sys.exit(1)
    with open(results_file) as f:
        results = json.load(f)
    generate_report(results)
